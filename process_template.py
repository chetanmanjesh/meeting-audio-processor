"""
Full pipeline: audio → transcript + MoM JSON + filled Gurkar & Associates docx.

Same as process.py, but additionally fills the template.docx with the extracted
MoM and saves it alongside the other outputs.

Usage:
    python process_template.py meeting.mp3
    python process_template.py part1.mp4 part2.mp4 part3.mp4 \\
        --subject "Indiranagar villa — review 1" \\
        --attendees "Priya, Rohit, Asha" \\
        --prepared-by "Studio team"
"""

import argparse
import copy
import json
import os
import shutil
import sys
import tempfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import process as base
import process_sarvam as sarvam_pipeline

DEFAULT_TEMPLATE = Path(__file__).parent / "template.docx"


# ---------- docx-writing helpers ----------

def set_cell_text(cell, text) -> None:
    """Replace a cell's text, preserving the formatting of its first run."""
    text = "" if text is None else str(text)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    if p.runs:
        first = p.runs[0]
        first.text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def set_cell_lead_and_detail(cell, lead: str, detail: str) -> None:
    """Set cell text as: **lead**: detail (lead bolded, detail plain)."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    # Wipe existing runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if lead:
        lead_run = p.add_run(lead)
        lead_run.bold = True
        if detail:
            p.add_run(f": {detail}")
    elif detail:
        p.add_run(detail)


def replace_paragraph_with_lines(paragraph, lines: list[str], bullet: bool = False) -> None:
    """Replace a placeholder paragraph with one or more lines."""
    if not lines:
        lines = ["—"]

    template_run = paragraph.runs[0] if paragraph.runs else None
    first_line = f"• {lines[0]}" if bullet else lines[0]
    if template_run:
        template_run.text = first_line
        for r in paragraph.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        paragraph.clear()
        paragraph.add_run(first_line)

    prev = paragraph
    for line in lines[1:]:
        text = f"• {line}" if bullet else line
        new_p = copy.deepcopy(paragraph._element)
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        prev._element.addnext(new_p)
        from docx.text.paragraph import Paragraph
        wrapper = Paragraph(new_p, paragraph._parent)
        wrapper.add_run(text)
        prev = wrapper


def replace_paragraph_with_structured(paragraph, blocks: list[dict]) -> None:
    """Replace a placeholder paragraph with structured content blocks.

    Each block dict is one of:
      {"kind": "heading", "text": "..."}     — bold sub-heading
      {"kind": "bullet", "lead": "...", "detail": "..."}  — '• **lead**: detail'
      {"kind": "bullet", "text": "..."}      — plain '• text'
      {"kind": "spacer"}                     — empty paragraph
    """
    from docx.text.paragraph import Paragraph

    if not blocks:
        blocks = [{"kind": "bullet", "text": "—"}]

    def apply_block(p, block):
        # Wipe existing runs first
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        kind = block.get("kind")
        if kind == "heading":
            run = p.add_run(block.get("text", ""))
            run.bold = True
        elif kind == "spacer":
            return
        else:
            p.add_run("• ")
            lead = block.get("lead", "")
            detail = block.get("detail", "")
            text = block.get("text", "")
            if lead:
                lead_run = p.add_run(lead)
                lead_run.bold = True
                if detail:
                    p.add_run(f": {detail}")
            elif detail:
                p.add_run(detail)
            elif text:
                p.add_run(text)

    apply_block(paragraph, blocks[0])

    prev = paragraph
    for block in blocks[1:]:
        new_p = copy.deepcopy(paragraph._element)
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        prev._element.addnext(new_p)
        wrapper = Paragraph(new_p, paragraph._parent)
        apply_block(wrapper, block)
        prev = wrapper


def find_paragraph_index(doc, predicate):
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i
    return None


def placeholder_after(doc, heading_pred):
    idx = find_paragraph_index(doc, heading_pred)
    if idx is None:
        return None
    paragraphs = doc.paragraphs
    for j in range(idx + 1, len(paragraphs)):
        if paragraphs[j].text.strip():
            return paragraphs[j]
    return None


def fill_meta_table(table, meta: dict) -> None:
    label_to_key = {
        "DATE": "date",
        "FORMAT": "format",
        "SUBJECT": "subject",
        "ATTENDEES": "attendees",
        "PREPARED BY": "prepared_by",
    }
    for row in table.rows:
        label = row.cells[0].text.strip().upper()
        key = label_to_key.get(label)
        if key:
            set_cell_text(row.cells[1], meta.get(key, "—"))


def fill_action_items_table(table, items: list[dict]) -> None:
    for row in list(table.rows[1:]):
        row._element.getparent().remove(row._element)

    if not items:
        new_row = table.add_row()
        for c in new_row.cells:
            set_cell_text(c, "—")
        return

    for item in items:
        new_row = table.add_row()
        set_cell_text(new_row.cells[0], item.get("owner") or "unassigned")
        set_cell_text(new_row.cells[1], item.get("task") or "")
        set_cell_text(new_row.cells[2], item.get("due") or "—")


def build_discussion_blocks(mom: dict) -> list[dict]:
    """Build the structured DISCUSSION NOTES content: sub-headings + bold-lead bullets.

    Sections in order:
      DECISIONS AGREED (themed groups of {lead, detail})
      SIGN-OFFS (flat, each with lead = space)
      MATERIALS & SPECS (flat, each with lead = element)
      DESIGN RATIONALE & DISCUSSION NOTES (themed groups of plain bullets)
      OPEN ITEMS rendered separately by main()
    """
    blocks: list[dict] = []

    def section_heading(text):
        blocks.append({"kind": "heading", "text": text})

    def sub_heading(text):
        blocks.append({"kind": "heading", "text": "   " + text})

    def bullet(lead=None, detail=None, text=None):
        blocks.append({"kind": "bullet", "lead": lead or "", "detail": detail or "", "text": text or ""})

    def spacer():
        blocks.append({"kind": "spacer"})

    if mom.get("decisions"):
        section_heading("DECISIONS AGREED")
        for group in mom["decisions"]:
            theme = group.get("theme", "General")
            sub_heading(theme)
            for it in group.get("items", []):
                bullet(lead=it.get("lead"), detail=it.get("detail"))
            spacer()

    if mom.get("sign_offs"):
        section_heading("SIGN-OFFS")
        for s in mom["sign_offs"]:
            lead = s.get("lead") or s.get("space_or_item", "—")
            tail = []
            if s.get("status"): tail.append(s["status"])
            if s.get("verbal_or_physical"): tail.append(s["verbal_or_physical"])
            if s.get("by_whom"): tail.append(f"by {s['by_whom']}")
            detail = ", ".join(tail)
            if s.get("conditions_or_notes"):
                detail = (detail + ". " if detail else "") + s["conditions_or_notes"]
            bullet(lead=lead, detail=detail)
        spacer()

    if mom.get("materials_and_specs"):
        section_heading("MATERIALS & SPECIFICATIONS")
        for m in mom["materials_and_specs"]:
            lead = m.get("space_or_element", "—")
            detail = m.get("material_or_spec", "")
            if m.get("notes"):
                detail = (detail + " — " if detail else "") + m["notes"]
            bullet(lead=lead, detail=detail)
        spacer()

    if mom.get("discussion_notes"):
        section_heading("DESIGN RATIONALE & DISCUSSION NOTES")
        for group in mom["discussion_notes"]:
            sub_heading(group.get("theme", "General"))
            for it in group.get("items", []):
                bullet(text=it)
            spacer()

    if mom.get("speakers"):
        section_heading("PARTICIPANTS")
        for s in mom["speakers"]:
            name = s.get("inferred_name") or s.get("label", "Speaker")
            role = f" — {s['role_guess']}" if s.get("role_guess") else ""
            sub_heading(name + role)
            for tp in s.get("talking_points", []):
                bullet(text=tp)
            spacer()

    return blocks


def build_open_items_blocks(mom: dict) -> list[dict]:
    """Render open_questions as 'TOPIC: pending description' bullets."""
    blocks: list[dict] = []
    for q in mom.get("open_questions", []) or []:
        topic = q.get("topic", "—") if isinstance(q, dict) else "—"
        pending = q.get("what_is_pending", "") if isinstance(q, dict) else str(q)
        blocks.append({"kind": "bullet", "lead": topic, "detail": pending})
    if not blocks:
        blocks.append({"kind": "bullet", "text": "—"})
    return blocks


def derive_attendees(mom: dict) -> str:
    speakers = mom.get("speakers") or []
    if not speakers:
        return "—"
    parts = []
    for s in speakers:
        name = s.get("inferred_name") or s.get("label", "Speaker")
        role = s.get("role_guess")
        parts.append(f"{name} ({role})" if role else name)
    return ", ".join(parts)


def fill_template_to_bytes(template_path: Path, mom: dict, meta: dict) -> bytes:
    """Same as fill_template but returns the docx as bytes (no disk write)."""
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        fill_template(template_path, mom, meta, tmp_path)
        return tmp_path.read_bytes()
    finally:
        tmp_path.unlink(missing_ok=True)


def auto_meta(mom: dict, prepared_by: str = "Studio team") -> dict:
    """Build default meta from the MoM itself — date today, format from meeting_mode,
    subject from summary, attendees from speakers."""
    summary = mom.get("summary") or "—"
    meeting_mode = mom.get("meeting_mode") or "Audio recording"
    subject = (summary.split(".")[0][:120] if summary and summary != "—" else "—")
    return {
        "date": date.today().strftime("%d %B %Y"),
        "format": meeting_mode,
        "subject": subject,
        "attendees": derive_attendees(mom) or "—",
        "prepared_by": prepared_by or "Studio team",
    }


def fill_template(template_path: Path, mom: dict, meta: dict, out_path: Path) -> None:
    doc = Document(str(template_path))

    fill_meta_table(doc.tables[0], meta)

    ph = placeholder_after(doc, lambda t: t.strip().upper() == "PURPOSE/SUMMARY")
    if ph is not None:
        replace_paragraph_with_lines(ph, [mom.get("summary", "—")])

    ph = placeholder_after(doc, lambda t: t.strip().upper() == "DISCUSSION NOTES")
    if ph is not None:
        replace_paragraph_with_structured(ph, build_discussion_blocks(mom))

    fill_action_items_table(doc.tables[1], mom.get("action_items", []))

    ph = placeholder_after(doc, lambda t: "OPEN ITEMS" in t.upper())
    if ph is not None:
        replace_paragraph_with_structured(ph, build_open_items_blocks(mom))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---------- main ----------

def main():
    parser = argparse.ArgumentParser(
        description="Audio → transcript + MoM JSON + filled Gurkar & Associates docx"
    )
    parser.add_argument(
        "audio", type=Path, nargs="+",
        help="Path(s) to audio file(s). Multiple files are stitched in order.",
    )
    parser.add_argument("--out", type=Path, default=Path("output"), help="Output directory")
    parser.add_argument(
        "--name", type=str, default=None,
        help="Output file stem. Defaults to first input's stem (with _combined if multiple).",
    )
    parser.add_argument(
        "--timeout", type=float, default=1800.0,
        help="Deepgram HTTP timeout in seconds (default 1800 = 30 min).",
    )
    parser.add_argument(
        "--provider", choices=["deepgram", "sarvam"], default="deepgram",
        help="ASR provider. Use 'sarvam' for meetings with Kannada/Hindi/Tamil etc. "
             "Deepgram (default) only handles English + Hindi cleanly.",
    )
    # Template metadata
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE,
                        help=f"Template docx (default: {DEFAULT_TEMPLATE.name})")
    parser.add_argument("--date", type=str, default=date.today().strftime("%d %B %Y"))
    parser.add_argument("--format", type=str, default="Audio recording", dest="meeting_format")
    parser.add_argument("--subject", type=str, default=None,
                        help="Defaults to first sentence of the summary")
    parser.add_argument("--attendees", type=str, default=None,
                        help="Comma-separated. Defaults to inferred speakers.")
    parser.add_argument("--prepared-by", type=str, default="Studio team")
    parser.add_argument("--no-concise", action="store_true",
                        help="Skip the concise second pass. Default: produce both _detailed and _concise docx/md/json.")
    args = parser.parse_args()

    for p in args.audio:
        if not p.exists():
            sys.exit(f"Audio file not found: {p}")

    asr_key_var = "DEEPGRAM_API_KEY" if args.provider == "deepgram" else "SARVAM_API_KEY"
    for var in (asr_key_var, "ANTHROPIC_API_KEY"):
        if not os.environ.get(var):
            sys.exit(f"Missing {var} in environment. Copy .env.example to .env and fill it in.")

    if not args.template.exists():
        sys.exit(f"Template not found: {args.template}")

    args.out.mkdir(parents=True, exist_ok=True)

    if args.name:
        stem = args.name
    elif len(args.audio) == 1:
        stem = args.audio[0].stem
    else:
        stem = f"{args.audio[0].stem}_combined"

    # Stitch (if needed), transcribe, extract MoM — reusing process.py's helpers.
    with tempfile.TemporaryDirectory() as tmp:
        if len(args.audio) > 1:
            stitched = base.stitch_audio(args.audio, Path(tmp) / stem)
            kept = args.out / f"{stem}{stitched.suffix}"
            shutil.copy2(stitched, kept)
            print(f"✓ Stitched:   {kept}", file=sys.stderr)
            audio_to_process = stitched
        else:
            audio_to_process = args.audio[0]

        if args.provider == "deepgram":
            asr_response = base.transcribe(audio_to_process, timeout_seconds=args.timeout)
            transcript = base.format_transcript(asr_response)
            asr_filename = f"{stem}.deepgram.json"
        else:
            asr_response = sarvam_pipeline.transcribe_sarvam(audio_to_process)
            transcript = sarvam_pipeline.format_sarvam_transcript(asr_response)
            asr_filename = f"{stem}.sarvam.json"

    (args.out / asr_filename).write_text(json.dumps(asr_response, indent=2, ensure_ascii=False))
    (args.out / f"{stem}.transcript.txt").write_text(transcript)
    print(f"✓ Transcript: {args.out / f'{stem}.transcript.txt'}", file=sys.stderr)

    # Detailed extraction (shared prompt — same for Deepgram and Sarvam paths)
    mom = base.extract_mom(transcript)
    (args.out / f"{stem}.mom_detailed.json").write_text(json.dumps(mom, indent=2, ensure_ascii=False))
    (args.out / f"{stem}.mom_detailed.md").write_text(base.render_markdown(mom))
    print(f"✓ Detailed (md): {args.out / f'{stem}.mom_detailed.md'}", file=sys.stderr)

    def build_meta(m: dict) -> dict:
        summary = m.get("summary", "—")
        return {
            "date": args.date,
            "format": m.get("meeting_mode") or args.meeting_format,
            "subject": args.subject or (summary.split(".")[0][:120] if summary else "—"),
            "attendees": args.attendees or derive_attendees(m),
            "prepared_by": args.prepared_by,
        }

    # Detailed docx
    detailed_docx = args.out / f"{stem}_detailed.docx"
    fill_template(args.template, mom, build_meta(mom), detailed_docx)
    print(f"✓ Detailed docx: {detailed_docx}", file=sys.stderr)

    # Concise version (no facts dropped, tighter prose)
    if not args.no_concise:
        concise = base.condense_to_concise(mom)
        (args.out / f"{stem}.mom_concise.json").write_text(json.dumps(concise, indent=2, ensure_ascii=False))
        (args.out / f"{stem}.mom_concise.md").write_text(base.render_markdown(concise))
        print(f"✓ Concise (md):  {args.out / f'{stem}.mom_concise.md'}", file=sys.stderr)

        concise_docx = args.out / f"{stem}_concise.docx"
        fill_template(args.template, concise, build_meta(concise), concise_docx)
        print(f"✓ Concise docx:  {concise_docx}", file=sys.stderr)


if __name__ == "__main__":
    main()
